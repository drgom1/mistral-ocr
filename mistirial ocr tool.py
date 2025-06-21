import tkinter as tk
from tkinter import ttk, filedialog, messagebox
import tkinterdnd2 as tkdnd
import requests
import base64
import threading
from pathlib import Path
import time
from docx import Document
import os
import subprocess
import platform
from collections import deque

class MistralOCRTool:
    COLORS = {
        'bg': '#0f172a', 'card': '#1e293b', 'primary': '#3b82f6',
        'hover': '#2563eb', 'text': '#e2e8f0', 'muted': '#94a3b8',
        'success': '#10b981', 'error': '#ef4444', 'input': '#334155',
        'link': '#60a5fa'
    }
    
    SUPPORTED_FORMATS = {
        '.pdf': 'application/pdf',
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg',
        '.png': 'image/png', '.gif': 'image/gif',
        '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation'
    }
    
    def __init__(self, root):
        self.root = root
        self.root.title("Mistral OCR")
        self.root.geometry("800x600")
        self.root.minsize(700, 500)
        self.root.configure(bg=self.COLORS['bg'])
        
        # Variables
        self.api_key = tk.StringVar()
        self.selected_files = []
        self.output_format = tk.StringVar(value="txt")
        self.include_images = tk.BooleanVar(value=True)
        self.image_limit = tk.IntVar(value=10)
        self.processed_outputs = deque(maxlen=10)  # Store last 10 outputs
        self.current_output_folder = None
        
        # Create menu bar
        self.create_menu()
        self.setup_ui()
        
    def create_menu(self):
        """Create application menu"""
        menubar = tk.Menu(self.root)
        self.root.config(menu=menubar)
        
        # File menu
        file_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="File", menu=file_menu)
        file_menu.add_command(label="Add Files", command=self.add_files, accelerator="Ctrl+O")
        file_menu.add_command(label="Clear Files", command=self.clear_files)
        file_menu.add_separator()
        file_menu.add_command(label="Open Output Folder", command=self.open_output_folder)
        file_menu.add_separator()
        file_menu.add_command(label="Exit", command=self.root.quit)
        
        # Edit menu
        edit_menu = tk.Menu(menubar, tearoff=0)
        menubar.add_cascade(label="Edit", menu=edit_menu)
        edit_menu.add_command(label="Clear Log", command=self.clear_log)
        
        # Bind shortcuts
        self.root.bind('<Control-o>', lambda e: self.add_files())
        
    def setup_ui(self):
        # Main container
        main = self.create_frame(self.root, self.COLORS['bg'], padx=20, pady=20)
        
        # Header
        self.create_header(main)
        
        # Content columns
        content = self.create_frame(main, self.COLORS['bg'], pady=(20, 0))
        left = self.create_frame(content, self.COLORS['bg'], side=tk.LEFT, padx=(0, 10))
        right = self.create_frame(content, self.COLORS['bg'], side=tk.RIGHT, padx=(10, 0))
        
        # Sections
        self.create_api_section(left)
        self.create_file_section(left)
        self.create_recent_outputs_section(left)
        self.create_options_section(right)
        self.create_log_section(right)
        self.create_process_section(main)
        
    def create_frame(self, parent, bg, **pack_opts):
        """Helper to create and pack frames"""
        frame = tk.Frame(parent, bg=bg)
        frame.pack(fill=tk.BOTH, expand=True, **pack_opts)
        return frame
        
    def create_card(self, parent):
        """Create card with consistent styling"""
        return self.create_frame(parent, self.COLORS['card'], pady=(0, 15), padx=15)
        
    def create_label(self, parent, text, **opts):
        """Helper for consistent labels"""
        defaults = {'bg': self.COLORS['card'], 'fg': self.COLORS['text'], 'font': ('Segoe UI', 11, 'bold')}
        defaults.update(opts)
        return tk.Label(parent, text=text, **defaults)
        
    def create_button(self, parent, text, command, style='default', **opts):
        """Helper for consistent buttons"""
        styles = {
            'default': {'bg': self.COLORS['input'], 'fg': self.COLORS['muted']},
            'primary': {'bg': self.COLORS['primary'], 'fg': 'white', 'font': ('Segoe UI', 12, 'bold')},
            'link': {'bg': self.COLORS['card'], 'fg': self.COLORS['link'], 'font': ('Segoe UI', 9, 'underline')}
        }
        config = {
            'relief': tk.FLAT, 'bd': 0, 'cursor': 'hand2',
            'activebackground': self.COLORS['hover'] if style != 'link' else self.COLORS['card']
        }
        config.update(styles.get(style, {}))
        config.update(opts)
        btn = tk.Button(parent, text=text, command=command, **config)
        
        # Add hover effect for link style
        if style == 'link':
            btn.bind("<Enter>", lambda e: btn.config(fg=self.COLORS['primary']))
            btn.bind("<Leave>", lambda e: btn.config(fg=self.COLORS['link']))
        
        return btn
        
    def create_header(self, parent):
        """Compact header with status"""
        header = self.create_frame(parent, self.COLORS['bg'], pady=0)
        
        # Title
        title_frame = self.create_frame(header, self.COLORS['bg'], side=tk.LEFT, pady=0)
        self.create_label(title_frame, "🔍 Mistral OCR", font=('Segoe UI', 20, 'bold')).pack(side=tk.LEFT)
        
        # Status dot
        self.status_dot = tk.Canvas(header, width=10, height=10, bg=self.COLORS['bg'], highlightthickness=0)
        self.status_dot.pack(side=tk.RIGHT, pady=10)
        self.status_id = self.status_dot.create_oval(0, 0, 10, 10, fill=self.COLORS['muted'], outline="")
        
    def create_api_section(self, parent):
        """API key input with toggle"""
        card = self.create_card(parent)
        self.create_label(card, "🔑 API Key").pack(anchor=tk.W, pady=(12, 8))
        
        input_frame = self.create_frame(card, self.COLORS['card'], pady=0)
        
        self.api_entry = tk.Entry(input_frame, textvariable=self.api_key, show="•",
                                 font=('Segoe UI', 10), bg=self.COLORS['input'],
                                 fg=self.COLORS['text'], relief=tk.FLAT, bd=8,
                                 insertbackground=self.COLORS['primary'])
        self.api_entry.pack(side=tk.LEFT, fill=tk.X, expand=True)
        
        self.show_btn = self.create_button(input_frame, "👁", self.toggle_api, padx=10)
        self.show_btn.pack(side=tk.RIGHT, padx=(8, 0))
        
    def create_file_section(self, parent):
        """File selection with drag & drop"""
        card = self.create_card(parent)
        
        # Header
        header = self.create_frame(card, self.COLORS['card'], pady=(12, 8))
        self.create_label(header, "📁 Files").pack(side=tk.LEFT)
        self.file_count = self.create_label(header, "0 selected", font=('Segoe UI', 9), fg=self.COLORS['muted'])
        self.file_count.pack(side=tk.RIGHT)
        
        # Drop area
        self.drop_area = tk.Frame(card, bg=self.COLORS['input'], height=60)
        self.drop_area.pack(fill=tk.BOTH, expand=True, pady=(0, 8))
        self.drop_area.pack_propagate(False)
        
        self.drop_label = self.create_label(self.drop_area, "📥 Drag files here\nor click to browse",
                                           bg=self.COLORS['input'], fg=self.COLORS['muted'],
                                           font=('Segoe UI', 10), cursor='hand2')
        self.drop_label.pack(expand=True)
        
        # Drag & drop setup
        self.drop_area.drop_target_register(tkdnd.DND_FILES)
        self.drop_area.dnd_bind('<<Drop>>', self.on_drop)
        self.drop_area.dnd_bind('<<DragEnter>>', lambda e: self.drop_area.configure(bg='#475569'))
        self.drop_area.dnd_bind('<<DragLeave>>', lambda e: self.drop_area.configure(bg=self.COLORS['input']))
        
        # Click handlers
        for widget in (self.drop_area, self.drop_label):
            widget.bind("<Button-1>", lambda e: self.add_files())
        
        # Clear button
        self.create_button(card, "Clear", self.clear_files, font=('Segoe UI', 9), 
                          padx=15, pady=5).pack(anchor=tk.E)
        
    def create_recent_outputs_section(self, parent):
        """Recent output files section"""
        card = self.create_card(parent)
        
        # Header
        header = self.create_frame(card, self.COLORS['card'], pady=(12, 8))
        self.create_label(header, "📄 Recent Outputs").pack(side=tk.LEFT)
        
        # Scrollable frame for outputs
        self.outputs_frame = tk.Frame(card, bg=self.COLORS['card'])
        self.outputs_frame.pack(fill=tk.BOTH, expand=True)
        
        # Initial message
        self.no_outputs_label = self.create_label(self.outputs_frame, 
                                                 "No outputs yet", 
                                                 font=('Segoe UI', 9), 
                                                 fg=self.COLORS['muted'])
        self.no_outputs_label.pack(pady=10)
        
    def create_options_section(self, parent):
        """Output options"""
        card = self.create_card(parent)
        self.create_label(card, "⚙️ Options").pack(anchor=tk.W, pady=(12, 8))
        
        # Format selection
        format_frame = self.create_frame(card, self.COLORS['card'], pady=(0, 8))
        self.create_label(format_frame, "Output:", font=('Segoe UI', 9), 
                         fg=self.COLORS['muted']).pack(side=tk.LEFT)
        
        for fmt, text in [("txt", "Text"), ("docx", "Word")]:
            tk.Radiobutton(format_frame, text=text, variable=self.output_format, value=fmt,
                          font=('Segoe UI', 9), bg=self.COLORS['card'], fg=self.COLORS['text'],
                          activebackground=self.COLORS['card'], selectcolor=self.COLORS['card']
                          ).pack(side=tk.LEFT, padx=(15, 0))
        
        # Image options
        img_frame = self.create_frame(card, self.COLORS['card'], pady=0)
        tk.Checkbutton(img_frame, text="Include images", variable=self.include_images,
                      font=('Segoe UI', 9), bg=self.COLORS['card'], fg=self.COLORS['text'],
                      activebackground=self.COLORS['card'], selectcolor=self.COLORS['card']
                      ).pack(side=tk.LEFT)
        
        self.create_label(img_frame, "Limit:", font=('Segoe UI', 9), 
                         fg=self.COLORS['muted']).pack(side=tk.LEFT, padx=(20, 5))
        
        # Validate spinbox input
        vcmd = (self.root.register(self.validate_number), '%P')
        tk.Spinbox(img_frame, from_=0, to=50, textvariable=self.image_limit, width=5,
                  font=('Segoe UI', 9), bg=self.COLORS['input'], fg=self.COLORS['text'],
                  buttonbackground='#475569', relief=tk.FLAT, bd=0,
                  validate='key', validatecommand=vcmd).pack(side=tk.LEFT)
        
    def create_log_section(self, parent):
        """Activity log with context menu"""
        card = self.create_card(parent)
        self.create_label(card, "📋 Activity").pack(anchor=tk.W, pady=(12, 8))
        
        log_frame = self.create_frame(card, self.COLORS['input'], pady=0)
        
        self.log = tk.Text(log_frame, height=10, font=('Consolas', 9),
                          bg=self.COLORS['input'], fg=self.COLORS['text'],
                          relief=tk.FLAT, bd=8, wrap=tk.WORD,
                          state=tk.DISABLED,  # Make read-only
                          insertbackground=self.COLORS['primary'])
        self.log.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        
        # Configure tags for clickable links
        self.log.tag_configure("link", foreground=self.COLORS['link'], underline=1)
        self.log.tag_bind("link", "<Button-1>", self.on_log_link_click)
        self.log.tag_bind("link", "<Enter>", lambda e: self.log.config(cursor="hand2"))
        self.log.tag_bind("link", "<Leave>", lambda e: self.log.config(cursor=""))
        
        scrollbar = ttk.Scrollbar(log_frame, command=self.log.yview)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)
        self.log.config(yscrollcommand=scrollbar.set)
        
        # Context menu
        self.log_menu = tk.Menu(self.log, tearoff=0)
        self.log_menu.add_command(label="Copy", command=self.copy_log_selection)
        self.log_menu.add_command(label="Clear Log", command=self.clear_log)
        self.log.bind("<Button-3>", self.show_log_menu)
        
        self.log_msg("Ready to process documents")
        
    def create_process_section(self, parent):
        """Process button with progress"""
        bottom = self.create_frame(parent, self.COLORS['bg'], pady=(10, 0))
        
        self.progress = ttk.Progressbar(bottom, mode='indeterminate')
        self.progress.pack(fill=tk.X, pady=(0, 10))
        
        # Button frame for multiple buttons
        btn_frame = self.create_frame(bottom, self.COLORS['bg'], pady=0)
        
        self.process_btn = self.create_button(btn_frame, "Process Documents", 
                                             self.process_docs, 'primary', padx=30, pady=12)
        self.process_btn.pack(side=tk.LEFT, padx=(0, 10))
        
        # Open folder button
        self.folder_btn = self.create_button(btn_frame, "📁 Output Folder", 
                                            self.open_output_folder, padx=20, pady=12)
        self.folder_btn.pack(side=tk.LEFT)
        self.folder_btn.config(state=tk.DISABLED)
        
    # Utility methods
    def validate_number(self, value):
        """Validate spinbox input"""
        if value == "":
            return True
        try:
            int(value)
            return True
        except ValueError:
            return False
    
    def open_file(self, filepath):
        """Cross-platform file opener"""
        if platform.system() == 'Windows':
            os.startfile(filepath)
        elif platform.system() == 'Darwin':  # macOS
            subprocess.call(['open', filepath])
        else:  # Linux
            subprocess.call(['xdg-open', filepath])
    
    def open_output_folder(self):
        """Open the folder containing the last output"""
        if self.current_output_folder and self.current_output_folder.exists():
            self.open_file(str(self.current_output_folder))
        else:
            messagebox.showinfo("Info", "No output folder available yet")
    
    def clear_log(self):
        """Clear the activity log"""
        self.log.config(state=tk.NORMAL)
        self.log.delete(1.0, tk.END)
        self.log.config(state=tk.DISABLED)
        self.log_msg("Log cleared")
    
    def copy_log_selection(self):
        """Copy selected text from log"""
        try:
            selection = self.log.get(tk.SEL_FIRST, tk.SEL_LAST)
            self.root.clipboard_clear()
            self.root.clipboard_append(selection)
        except tk.TclError:
            pass  # No selection
    
    def show_log_menu(self, event):
        """Show context menu for log"""
        self.log_menu.post(event.x_root, event.y_root)
    
    def on_log_link_click(self, event):
        """Handle clicks on file links in log"""
        # Get the index of the click
        index = self.log.index(f"@{event.x},{event.y}")
        # Get the file path from the tag
        tags = self.log.tag_names(index)
        for tag in tags:
            if tag.startswith("file:"):
                filepath = tag[5:]  # Remove "file:" prefix
                if Path(filepath).exists():
                    self.open_file(filepath)
                break
    
    def update_recent_outputs(self):
        """Update the recent outputs display"""
        # Clear existing widgets
        for widget in self.outputs_frame.winfo_children():
            widget.destroy()
        
        if not self.processed_outputs:
            self.no_outputs_label = self.create_label(self.outputs_frame, 
                                                     "No outputs yet", 
                                                     font=('Segoe UI', 9), 
                                                     fg=self.COLORS['muted'])
            self.no_outputs_label.pack(pady=10)
        else:
            for output_path in self.processed_outputs:
                output_frame = tk.Frame(self.outputs_frame, bg=self.COLORS['card'])
                output_frame.pack(fill=tk.X, pady=2)
                
                # File icon and name
                filename = Path(output_path).name
                display_name = filename if len(filename) <= 30 else filename[:27] + "..."
                
                file_btn = self.create_button(output_frame, f"📄 {display_name}", 
                                            lambda p=output_path: self.open_file(p), 
                                            style='link', anchor='w')
                file_btn.pack(side=tk.LEFT, fill=tk.X, expand=True)
                
                # Open folder button
                folder_btn = self.create_button(output_frame, "📁", 
                                              lambda p=output_path: self.open_file(str(Path(p).parent)),
                                              font=('Segoe UI', 8), padx=5)
                folder_btn.pack(side=tk.RIGHT)
    
    # Functionality methods
    def toggle_api(self):
        """Toggle password visibility"""
        current = self.api_entry.cget('show')
        self.api_entry.config(show='' if current else '•')
        self.show_btn.config(text='🙈' if current else '👁')
    
    def on_drop(self, event):
        """Handle file drop"""
        try:
            files = self.root.tk.splitlist(event.data)
            # Filter out directories and validate files
            valid_files = []
            for f in files:
                path = Path(f)
                if path.is_file() and path.suffix.lower() in self.SUPPORTED_FORMATS:
                    valid_files.append(f)
                elif path.is_dir():
                    self.log_msg(f"⚠️ Skipped directory: {path.name}")
            
            if valid_files:
                self.add_files_list(valid_files)
            elif not any(Path(f).is_dir() for f in files):
                self.log_msg("❌ No supported files found")
        except Exception as e:
            self.log_msg(f"❌ Drop error: {str(e)}")
        finally:
            self.drop_area.configure(bg=self.COLORS['input'])
    
    def add_files(self):
        """Browse for files"""
        extensions = ";".join(f"*{ext}" for ext in self.SUPPORTED_FORMATS)
        files = filedialog.askopenfilenames(
            title="Select files to process",
            filetypes=[("Supported Files", extensions), ("All Files", "*.*")]
        )
        if files:
            self.add_files_list(files)
    
    def add_files_list(self, files):
        """Add files to selection"""
        new_files = [f for f in files if f not in self.selected_files]
        self.selected_files.extend(new_files)
        self.update_file_count()
        if new_files:
            self.log_msg(f"Added {len(new_files)} file(s)")
    
    def clear_files(self):
        """Clear file selection"""
        self.selected_files.clear()
        self.update_file_count()
        self.log_msg("Cleared all files")
    
    def update_file_count(self):
        """Update file count display"""
        count = len(self.selected_files)
        self.file_count.config(text=f"{count} selected")
        text = "📥 Drag more files\nor click to browse" if count else "📥 Drag files here\nor click to browse"
        self.drop_label.config(text=text)
    
    def log_msg(self, msg, file_path=None):
        """Thread-safe logging with optional file link"""
        def _log():
            self.log.config(state=tk.NORMAL)
            timestamp = f"[{time.strftime('%H:%M')}] "
            self.log.insert(tk.END, timestamp)
            
            if file_path and Path(file_path).exists():
                # Insert message with file link
                parts = msg.split(Path(file_path).name)
                if len(parts) > 1:
                    self.log.insert(tk.END, parts[0])
                    # Insert filename as clickable link
                    start = self.log.index(tk.END + "-1c")
                    self.log.insert(tk.END, Path(file_path).name)
                    end = self.log.index(tk.END + "-1c")
                    self.log.tag_add("link", start, end)
                    self.log.tag_add(f"file:{file_path}", start, end)
                    self.log.insert(tk.END, parts[1])
                else:
                    self.log.insert(tk.END, msg)
            else:
                self.log.insert(tk.END, msg)
            
            self.log.insert(tk.END, "\n")
            self.log.see(tk.END)
            self.log.config(state=tk.DISABLED)
            
        self.root.after(0, _log)
    
    def update_status(self, color):
        """Update status indicator"""
        self.root.after(0, lambda: self.status_dot.itemconfig(self.status_id, fill=color))
    
    def process_docs(self):
        """Validate and start processing"""
        if not self.api_key.get().strip():
            messagebox.showerror("Error", "Please enter API key")
            return
        
        if not self.selected_files:
            messagebox.showerror("Error", "Please select files")
            return
        
        threading.Thread(target=self._process_thread, daemon=True).start()
    
    def _process_thread(self):
        """Processing thread"""
        def update_ui(state, text, progress_active, folder_enabled=False):
            self.process_btn.config(state=state, text=text)
            self.folder_btn.config(state=tk.NORMAL if folder_enabled else tk.DISABLED)
            if progress_active:
                self.progress.start(10)
            else:
                self.progress.stop()
        
        self.root.after(0, lambda: update_ui('disabled', 'Processing...', True))
        self.update_status(self.COLORS['primary'])
        
        try:
            success_count = 0
            for i, file in enumerate(self.selected_files, 1):
                self.log_msg(f"Processing {i}/{len(self.selected_files)}: {Path(file).name}")
                if self._process_file(file):
                    success_count += 1
            
            if success_count > 0:
                self.log_msg(f"✅ Processing complete! ({success_count}/{len(self.selected_files)} successful)")
                self.update_status(self.COLORS['success'])
                self.root.after(0, self.update_recent_outputs)
                self.root.after(0, lambda: update_ui('normal', 'Process Documents', False, True))
            else:
                self.log_msg("❌ No files were processed successfully")
                self.update_status(self.COLORS['error'])
                self.root.after(0, lambda: update_ui('normal', 'Process Documents', False, False))
            
        except Exception as e:
            self.log_msg(f"❌ Error: {str(e)}")
            self.update_status(self.COLORS['error'])
            self.root.after(0, lambda: update_ui('normal', 'Process Documents', False))
            
        finally:
            time.sleep(2)
            self.update_status(self.COLORS['muted'])
    
    def _process_file(self, file_path):
        """Process single file - returns True if successful"""
        try:
            # Validate file size
            file_size = Path(file_path).stat().st_size
            if file_size > 100 * 1024 * 1024:  # 100MB limit
                self.log_msg(f"⚠️ File too large: {Path(file_path).name} ({file_size // 1024 // 1024}MB)")
                return False
            
            # Read and encode file
            with open(file_path, "rb") as f:
                encoded = base64.b64encode(f.read()).decode('utf-8')
            
            ext = Path(file_path).suffix.lower()
            mime = self.SUPPORTED_FORMATS.get(ext, 'application/octet-stream')
            
            # API call
            response = requests.post(
                "https://api.mistral.ai/v1/ocr",
                headers={
                    "Authorization": f"Bearer {self.api_key.get()}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "mistral-ocr-latest",
                    "document": {
                        "type": "document_url",
                        "document_url": f"data:{mime};base64,{encoded}"
                    },
                    "include_image_base64": self.include_images.get(),
                    "image_limit": self.image_limit.get()
                },
                timeout=300
            )
            
            if response.status_code == 200:
                return self._save_results(file_path, response.json())
            else:
                self.log_msg(f"API Error {response.status_code}: {response.text[:100]}")
                return False
                
        except requests.exceptions.Timeout:
            self.log_msg("Request timeout - file may be too large")
            return False
        except requests.exceptions.RequestException as e:
            self.log_msg(f"Network error: {str(e)}")
            return False
        except Exception as e:
            self.log_msg(f"Failed to process {Path(file_path).name}: {str(e)}")
            return False
    
    def _save_results(self, file_path, result):
        """Save OCR results - returns True if successful"""
        try:
            base_path = Path(file_path)
            pages = result.get('pages', [])
            
            if not pages:
                self.log_msg("No content found in response")
                return False
            
            # Determine output path
            output_dir = base_path.parent
            base_name = base_path.stem
            extension = ".docx" if self.output_format.get() == "docx" else ".txt"
            
            # Find unique filename
            output = output_dir / f"{base_name}_ocr{extension}"
            counter = 1
            while output.exists():
                output = output_dir / f"{base_name}_ocr_{counter}{extension}"
                counter += 1
            
            if self.output_format.get() == "docx":
                # Save as DOCX
                doc = Document()
                doc.add_heading(f'OCR Results - {base_path.name}', 0)
                
                for page in pages:
                    doc.add_heading(f'Page {page.get("index", "?")}', 1)
                    content = page.get('markdown', '')
                    if content:
                        doc.add_paragraph(content)
                
                doc.save(str(output))
            else:
                # Save as TXT
                with open(output, 'w', encoding='utf-8') as f:
                    f.write(f"OCR Results - {base_path.name}\n")
                    f.write("=" * 50 + "\n\n")
                    
                    for page in pages:
                        f.write(f"=== Page {page.get('index', '?')} ===\n")
                        content = page.get('markdown', '')
                        if content:
                            f.write(content + "\n\n")
            
            # Update tracking
            self.processed_outputs.append(str(output))
            self.current_output_folder = output.parent
            
            # Log with clickable link
            self.log_msg(f"✓ Saved: {output.name}", str(output))
            return True
            
        except PermissionError:
            self.log_msg(f"❌ Permission denied: Cannot save to {output_dir}")
            return False
        except Exception as e:
            self.log_msg(f"❌ Save error: {str(e)}")
            return False

def main():
    root = tkdnd.TkinterDnD.Tk()
    MistralOCRTool(root)
    root.mainloop()

if __name__ == "__main__":
    main()